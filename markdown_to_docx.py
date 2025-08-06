import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from bs4 import BeautifulSoup

def markdown_to_docx_with_images(markdown_file_path, output_dir=None):
    """
    将markdown文件转换为docx文件，并处理图片（支持Markdown和HTML格式）
    
    Args:
        markdown_file_path (str): markdown文件路径
        output_dir (str, optional): 输出目录，默认为markdown文件所在目录
    
    Returns:
        str: 生成的docx文件路径
    """
    markdown_path = Path(markdown_file_path)
    
    # 检查文件是否存在
    if not markdown_path.exists():
        raise FileNotFoundError(f"Markdown文件不存在: {markdown_file_path}")
    
    # 设置输出目录和文件名
    if output_dir is None:
        output_dir = markdown_path.parent
    else:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
    
    docx_file_path = output_dir / f"{markdown_path.stem}.docx"
    
    # 读取markdown文件
    with open(markdown_path, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置文档样式
    setup_document_styles(doc)
    
    # 解析markdown内容
    lines = markdown_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 处理空行
        if not line:
            i += 1
            continue
        print(line)
        # 处理HTML格式的图片（包括div包装的图片）
        if '<img' in line or '<div' in line:
            print("befor:",i)
            i = process_html_content(doc, lines, i, markdown_path.parent)
            print("after:",i)
            continue
        
        # 处理标题
        elif line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title_text = line.lstrip('#').strip()
            if title_text:
                heading = doc.add_heading(title_text, level=min(level, 6))
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 处理Markdown格式的图片
        elif '![' in line and '](' in line:
            process_markdown_image_line(doc, line, markdown_path.parent)
        
        # 处理代码块
        elif line.startswith('```'):
            i = process_code_block(doc, lines, i)
            continue
        
        # 处理表格
        elif '|' in line and line.count('|') >= 2:
            i = process_table(doc, lines, i)
            continue
        
        # 处理列表
        elif line.startswith(('- ', '* ', '+ ')) or re.match(r'^\d+\. ', line):
            i = process_list(doc, lines, i)
            continue
        
        # 处理引用
        elif line.startswith('>'):
            process_quote(doc, line)
        
        # 处理普通段落
        else:
            process_paragraph(doc, line)
        
        i += 1
    
    # 保存文档
    doc.save(docx_file_path)
    print(f"转换完成: {docx_file_path}")
    return str(docx_file_path)

def process_html_content(doc, lines, start_index, base_path):
    """
    处理HTML内容，特别是图片标签
    """
    i = start_index
    html_content = ""
    max_lines = 10  # 添加最大行数限制，防止死循环
    lines_processed = 0
    
    # 收集HTML内容（可能跨多行）
    while i < len(lines) and lines_processed < max_lines:
        line = lines[i]
        html_content += line + "\n"
        lines_processed += 1
        i += 1
        # 如果是单行HTML或者遇到闭合标签，停止收集
        if ('<img' in line and '/>' in line) or \
           ('</div>' in line) or \
           ('</p>' in line) or \
           ('</table>' in line) or \
           ('</html>' in line) or \
           (not line.strip().startswith('<') and html_content.strip() and lines_processed > 1):
            break

    # 如果达到最大行数限制，记录警告
    if lines_processed >= max_lines:
        print(f"警告：HTML内容处理达到最大行数限制({max_lines}行)，可能存在格式问题")
    
    # 解析HTML内容
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        print("html_content:",html_content)
        # 查找所有图片标签
        img_tags = soup.find_all('img')
        
        for img_tag in img_tags:
            src = img_tag.get('src', '')
            alt = img_tag.get('alt', '')
            width = img_tag.get('width', '')
            
            if src:
                process_image_from_path(doc, src, alt, base_path, width)
        
        # 处理表格内容
        tables = soup.find_all('table')
        for table in tables:
            process_html_table(doc, table)
        
        # 处理其他文本内容
        text_content = soup.get_text().strip()
        if text_content and not img_tags and not tables:
            doc.add_paragraph(text_content)
            
    except Exception as e:
        print(f"解析HTML内容时出错: {e}")
        # 如果解析失败，作为普通文本处理
        doc.add_paragraph(html_content.strip())
    
    return i

def process_html_table(doc, table_soup):
    """
    处理HTML表格
    """
    try:
        rows = table_soup.find_all('tr')
        if not rows:
            return
        
        # 计算最大列数
        max_cols = 0
        for row in rows:
            cols = len(row.find_all(['td', 'th']))
            max_cols = max(max_cols, cols)
        
        if max_cols == 0:
            return
        
        # 创建Word表格
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        # 填充表格数据
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for col_idx, cell in enumerate(cells):
                if col_idx < max_cols:
                    table.cell(row_idx, col_idx).text = cell.get_text().strip()
                    
    except Exception as e:
        print(f"处理HTML表格时出错: {e}")

def process_image_from_path(doc, img_path, alt_text, base_path, width_attr=None):
    """
    从路径处理图片并添加到文档
    """
    try:
        # 处理相对路径
        if not os.path.isabs(img_path):
            full_img_path = base_path / img_path
        else:
            full_img_path = Path(img_path)
        
        # 检查图片文件是否存在
        if full_img_path.exists():
            # 添加图片到文档
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
            
            # 获取图片尺寸并调整
            try:
                with Image.open(full_img_path) as img:
                    width, height = img.size
                    
                    # 处理width属性（如果有的话）
                    if width_attr:
                        if width_attr.endswith('%'):
                            # 百分比宽度，转换为英寸（假设页面宽度为6.5英寸）
                            percent = float(width_attr.rstrip('%')) / 100
                            new_width = Inches(6.5 * percent)
                        else:
                            # 像素宽度，转换为英寸
                            new_width = Inches(float(width_attr) / 100)
                    else:
                        # 默认处理：设置最大宽度为6英寸
                        max_width = Inches(6)
                        if width > height:
                            new_width = min(max_width, Inches(width/100))
                        else:
                            new_height = min(Inches(4), Inches(height/100))
                            new_width = Inches(width * new_height.inches / (height/100))
                    
                    # 计算对应的高度
                    aspect_ratio = height / width
                    new_height = Inches(new_width.inches * aspect_ratio)
                
                run = paragraph.add_run()
                run.add_picture(str(full_img_path), width=new_width, height=new_height)
                
                # 添加图片说明
                if alt_text and alt_text.lower() != 'image':
                    caption_para = doc.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.add_run(f"图: {alt_text}")
                    caption_run.font.size = Pt(9)
                    caption_run.font.italic = True
                    
            except Exception as e:
                print(f"处理图片时出错 {full_img_path}: {e}")
                # 如果图片处理失败，添加文本说明
                doc.add_paragraph(f"[图片: {alt_text or img_path}]")
        else:
            print(f"图片文件不存在: {full_img_path}")
            doc.add_paragraph(f"[图片不存在: {alt_text or img_path}]")
            
    except Exception as e:
        print(f"处理图片时出错: {e}")
        doc.add_paragraph(f"[图片处理错误: {alt_text or img_path}]")

def process_markdown_image_line(doc, line, base_path):
    """
    处理Markdown格式的图片行
    """
    # 匹配图片语法: ![alt](path)
    img_pattern = r'!\[([^\]]*)\]\(([^\)]+)\)'
    matches = re.findall(img_pattern, line)
    
    for alt_text, img_path in matches:
        process_image_from_path(doc, img_path, alt_text, base_path)

def setup_document_styles(doc):
    """设置文档样式"""
    # 设置正文样式
    styles = doc.styles
    if 'Normal' in styles:
        normal_style = styles['Normal']
        normal_style.font.name = '微软雅黑'
        normal_style.font.size = Pt(11)

def process_code_block(doc, lines, start_index):
    """处理代码块"""
    i = start_index + 1
    code_lines = []
    
    # 收集代码块内容
    while i < len(lines) and not lines[i].startswith('```'):
        code_lines.append(lines[i])
        i += 1
    
    # 添加代码块到文档
    if code_lines:
        code_text = '\n'.join(code_lines)
        para = doc.add_paragraph()
        run = para.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        # 设置代码块样式
        para.style = doc.styles['Normal']
    
    return i

def process_table(doc, lines, start_index):
    """处理表格"""
    table_lines = []
    i = start_index
    
    # 收集表格行
    while i < len(lines) and '|' in lines[i]:
        line = lines[i].strip()
        if line and not re.match(r'^\|?\s*:?-+:?\s*\|', line):  # 跳过分隔行
            table_lines.append(line)
        i += 1
    
    if len(table_lines) >= 1:
        # 解析表格数据
        rows = []
        for line in table_lines:
            cells = [cell.strip() for cell in line.split('|')]
            # 移除首尾空元素
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]
            if cells:
                rows.append(cells)
        
        if rows:
            # 创建表格
            max_cols = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = 'Table Grid'
            
            # 填充表格数据
            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < max_cols:
                        table.cell(row_idx, col_idx).text = cell_data
    
    return i - 1

def process_list(doc, lines, start_index):
    """处理列表"""
    i = start_index
    list_items = []
    
    # 收集列表项
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        
        if line.startswith(('- ', '* ', '+ ')):
            list_items.append(('bullet', line[2:].strip()))
        elif re.match(r'^\d+\. ', line):
            list_items.append(('number', re.sub(r'^\d+\. ', '', line).strip()))
        else:
            break
        i += 1
    
    # 添加列表到文档
    for list_type, item_text in list_items:
        para = doc.add_paragraph()
        if list_type == 'bullet':
            para.style = 'List Bullet'
        else:
            para.style = 'List Number'
        para.add_run(item_text)
    
    return i - 1

def process_quote(doc, line):
    """处理引用"""
    quote_text = line.lstrip('>').strip()
    if quote_text:
        para = doc.add_paragraph()
        run = para.add_run(quote_text)
        run.font.italic = True
        para.style = 'Quote'

def process_paragraph(doc, line):
    """处理普通段落"""
    if line.strip():
        # 处理内联格式
        para = doc.add_paragraph()
        process_inline_formatting(para, line)

def process_inline_formatting(paragraph, text):
    """处理内联格式（粗体、斜体等）"""
    # 简单的内联格式处理
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 粗体
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # 斜体
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # 代码
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            # 普通文本
            paragraph.add_run(part)

# 使用示例
if __name__ == "__main__":
    # 示例用法
    markdown_file = "output/北邮~结构图测试2.md"  # 替换为您的markdown文件路径
    
    try:
        docx_path = markdown_to_docx_with_images(markdown_file)
        print(f"转换成功！生成的Word文档: {docx_path}")
    except Exception as e:
        print(f"转换失败: {e}")
